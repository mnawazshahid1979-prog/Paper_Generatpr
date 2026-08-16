import os
import re
import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image as PILImage
import matplotlib.pyplot as plt
import io

# صفحہ کی سیٹنگ
st.set_page_config(
    page_title="Government High School Khan Pur Maral - Paper Generator",
    page_icon="📝",
    layout="centered"
)

# شاندار CSS اسٹائلنگ
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/earlyaccess/jameelnoorinastaleeq.css');
    
    .stApp, .main, [data-testid="stAppViewContainer"] {
        background-color: #f3f4f6 !important; 
    }

    .stSelectbox, .stTextInput, div[data-baseweb="select"] {
        font-family: 'Jameel Noori Nastaleeq', sans-serif !important;
        font-size: 14px !important;
        direction: ltr !important;
        text-align: left !important;
    }
    p, label, span, div, h3, h4 {
        font-family: 'Jameel Noori Nastaleeq', sans-serif !important;
        font-size: 14px !important;
        direction: ltr !important;
        text-align: left !important;
    }
    
    div.stButton > button {
        background: linear-gradient(135deg, #059669 100%, #047857 0%);
        color: white;
        font-weight: bold;
        font-size: 14px;
        font-family: 'Jameel Noori Nastaleeq', sans-serif !important;
        border-radius: 8px;
        padding: 0px;
        width: 2.0in !important;
        height: 0.5in !important;
        border: none;
        box-shadow: 0 4px 12px rgba(5, 150, 105, 0.35);
        transition: 0.3s;
        display: block;
        margin-top: -15px !important;
        margin-left: 0px !important;
        margin-right: auto !important;
    }
    div.stButton > button:hover {
        background: linear-gradient(135deg, #047857 100%, #065f46 0%);
        box-shadow: 0 6px 16px rgba(5, 150, 105, 0.45);
    }
    .unique-question-card {
        background: #ffffff;
        border-right: 5px solid #2563eb;
        border-top: 1px solid #e2e8f0;
        border-left: 1px solid #e2e8f0;
        border-bottom: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 16px 20px;
        margin-bottom: 14px;
        box-shadow: 0 4px 15px rgba(37, 99, 235, 0.05);
        transition: all 0.3s ease-in-out;
    }
    .unique-question-card:hover {
        box-shadow: 0 6px 20px rgba(37, 99, 235, 0.1);
        transform: translateY(-2px);
    }
    .header-container {
        display: flex;
        justify-content: center;
        align-items: center;
        width: 100%;
        margin-bottom: 10px;
    }
    .header-box-2 {
        background: linear-gradient(135deg, #c0504d 0%, #9b3b38 100%);
        padding: 12px 30px;
        border-radius: 10px;
        color: white;
        box-shadow: 0 4px 12px rgba(192, 80, 77, 0.25);
        font-family: 'Jameel Noori Nastaleeq', Arial, sans-serif !important;
        font-weight: 700;
        font-size: 32px !important;
        letter-spacing: 1px;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        display: inline-block;
        text-align: center;
    }
    .header-box-3 {
        background: transparent !important;
        padding: 5px;
        text-align: center !important;
        color: #1e293b;
        margin-bottom: 20px;
        font-family: 'Arial', sans-serif !important;
        font-size: 18px !important;
        font-weight: bold;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <div class="header-container">
        <div class="header-box-2">Paper-Generator-Software</div>
    </div>
    <div class="header-box-3">Developed by M Nawaz Shahid M.Sc Math</div>
    """,
    unsafe_allow_html=True
)

excel_path = "Question_Bank.xlsx"

# ==============================================================================
# [ 🎯 ] MATPLOTLIB BASED MATH EXPRESSION TO IMAGE GENERATOR (Clear & Readable)
# ==============================================================================
def add_math_expression_as_image(paragraph, latex_expr, font_size=28):
    if not latex_expr or not str(latex_expr).strip():
        return False
    
    latex_expr = str(latex_expr).strip()
    clean_expr = latex_expr.replace("$", "").strip()
    
    try:
        fig, ax = plt.subplots(figsize=(6, 1.2), dpi=600)
        ax.axis('off')
        ax.text(0.5, 0.5, f"${clean_expr}$", fontsize=font_size, ha='center', va='center', 
                color='black', fontfamily='serif', fontweight='bold')
        
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', bbox_inches='tight', transparent=True, dpi=600)
        plt.close(fig)
        img_stream.seek(0)
        
        run = paragraph.add_run()
        run.add_picture(img_stream, height=Pt(24))
        return True
    except Exception:
        paragraph.add_run(latex_expr)
        return True

def load_excel_data(path):
    if not os.path.exists(path):
        return None
    try:
        xls = pd.ExcelFile(path)
        sheets = xls.sheet_names
        data_structure = {}

        for sheet in sheets:
            if "_" in sheet:
                cls, subj = sheet.split("_", 1)
                cls_name = cls + " Class"
            else:
                cls_name = "Other"
                subj = sheet

            df = pd.read_excel(path, sheet_name=sheet)
            units_data = {}
            if "Unit" in df.columns and "Topic_No" in df.columns:
                for unit_val in df["Unit"].dropna().unique():
                    topics = (
                        df[df["Unit"] == unit_val]["Topic_No"]
                        .dropna()
                        .astype(str)
                        .str.strip()
                        .unique()
                        .tolist()
                    )
                    units_data[str(unit_val).strip()] = topics

            if cls_name not in data_structure:
                data_structure[cls_name] = {}

            data_structure[cls_name][subj] = {
                "sheet": sheet,
                "units_data": units_data,
            }
        return data_structure
    except Exception as e:
        st.error(f"اكسل فائل پڑھنے میں مسئلہ پیش آیا:\n{e}")
        return None

data_structure = load_excel_data(excel_path)

if data_structure is None:
    st.error(f"مطلوبہ ایکسل فائل '{excel_path}' فولڈر میں نہیں ملی!")
else:
    if "saved_sections" not in st.session_state:
        st.session_state.saved_sections = set()
    if "section_data_records" not in st.session_state:
        st.session_state.section_data_records = {}
    if "start_process" not in st.session_state:
        st.session_state.start_process = False
    if "show_questions_flag" not in st.session_state:
        st.session_state.show_questions_flag = False
    if "selected_section_tab" not in st.session_state:
        st.session_state.selected_section_tab = "MCQs"
    if not st.session_state.start_process:
        col_btn, _ = st.columns([1, 4])
        with col_btn:
            if st.button("پیپر منتخب کریں"):
                st.session_state.start_process = True
                st.rerun()
    else:
        st.markdown("### 🏛️ ادارے کی معلومات اور لوگو")
        
        institution_name = st.text_input(
            "اپنے ادارے کا نام لکھیں (لکھ کر Enter دبائیں):", 
            value="", 
            placeholder="یہاں اپنے اسکول یا ادارے کا نام لکھیں..."
        )        
        uploaded_logo = None
        proceed_to_paper = False

        if institution_name.strip() != "":
            logo_choice = st.selectbox("کیا آپ لوگو شامل کرنا چاہتے ہیں؟", ["", "No", "Yes"], key="logo_choice_dropdown")
            
            if logo_choice == "No":
                proceed_to_paper = True
            elif logo_choice == "Yes":
                uploaded_logo = st.file_uploader("لوگو اپ لوڈ کریں", type=["png", "jpg", "jpeg"])
                if uploaded_logo is not None:
                    proceed_to_paper = True

        st.markdown("---")

        if proceed_to_paper:
            classes = list(data_structure.keys())
            selected_cls = st.selectbox("1. کلاس منتخب کریں:", [""] + classes, key="sel_cls")

            selected_subj = ""
            if selected_cls != "":
                subj_list = list(data_structure[selected_cls].keys())
                selected_subj = st.selectbox("2. مضمون منتخب کریں:", [""] + subj_list, key="sel_subj")

            selected_mode = ""
            if selected_subj != "":
                selected_mode = st.selectbox(
                    "3. پیپر موڈ منتخب کریں:",
                    ["", "سنگل یونٹ (Single Unit)", "ملٹی یونٹ (Multi Unit)", "فل بک (Full Book)", "ٹاپک وائز (Topic Wise)"],
                    key="sel_mode"
                )

            selected_unit = ""
            selected_topics = []
            
            if selected_mode != "" and "ٹاپک وائز" in selected_mode:
                units_data = data_structure[selected_cls][selected_subj]["units_data"]
                unit_str_list = [str(u) for u in units_data.keys()]
                selected_unit = st.selectbox("4. یونٹ منتخب کریں:", [""] + unit_str_list, key="sel_unit")

                if selected_unit != "":
                    topics = units_data.get(str(selected_unit).strip(), [])
                    if topics:
                        st.markdown("##### 📌 ٹاپکس منتخب کریں:")
                        for topic in topics:
                            t_key = f"top_{selected_cls}_{selected_subj}_{selected_unit}_{topic}"
                            is_checked = st.checkbox(str(topic), key=t_key)
                            if is_checked:
                                selected_topics.append(topic)

            can_update = False
            if selected_cls != "" and selected_subj != "" and selected_mode != "":
                if "ٹاپک وائز" in selected_mode:
                    if selected_unit != "" and len(selected_topics) > 0:
                        can_update = True
                else:
                    can_update = True

            if can_update:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("اپ ڈیٹ", key="update_btn_action"):
                    st.session_state.show_questions_flag = True
                    st.rerun()

            st.markdown("---")

            if st.session_state.show_questions_flag and can_update:
                st.markdown("### 📚 سوالات کے سیکشنز اور انتخاب")

                col_t1, col_t2, col_t3 = st.columns(3)
                
                with col_t1:
                    is_mcq_saved = 1 in st.session_state.saved_sections
                    btn_label_1 = "🔒 معروضی (MCQs) - محفوظ" if is_mcq_saved else "📝 معروضی (MCQs)"
                    if st.button(btn_label_1, key="btn_tab_mcq", disabled=is_mcq_saved):
                        st.session_state.selected_section_tab = "MCQs"
                        st.rerun()

                with col_t2:
                    is_sq_saved = 2 in st.session_state.saved_sections
                    btn_label_2 = "🔒 مختصر (SQs) - محفوظ" if is_sq_saved else "📝 مختصر سوالات (SQs)"
                    if st.button(btn_label_2, key="btn_tab_sq", disabled=is_sq_saved):
                        st.session_state.selected_section_tab = "SQs"
                        st.rerun()

                with col_t3:
                    is_lq_saved = 5 in st.session_state.saved_sections
                    btn_label_3 = "🔒 طویل (LQs) - محفوظ" if is_lq_saved else "📝 طویل سوالات (LQs)"
                    if st.button(btn_label_3, key="btn_tab_lq", disabled=is_lq_saved):
                        st.session_state.selected_section_tab = "LQs"
                        st.rerun()

                st.markdown("---")

                sheet_info = data_structure[selected_cls][selected_subj]
                actual_sheet = sheet_info["sheet"]

                try:
                    df_full = pd.read_excel(excel_path, sheet_name=actual_sheet)

                    current_tab = st.session_state.selected_section_tab
                    if current_tab == "MCQs":
                        q_id, title_text = 1, "حصہ اول - معروضی سوالات"
                    elif current_tab == "SQs":
                        q_id, title_text = 2, "حصہ دوم - مختصر سوالات"
                    else:
                        q_id, title_text = 5, "حصہ سوم - طویل سوالات"

                    st.markdown(f"#### **{title_text}**")
                    filtered_df = df_full.copy()
                    if "Q_Type" in filtered_df.columns:
                        filtered_df["Q_Type"] = filtered_df["Q_Type"].astype(str).str.strip().str.upper()
                        if q_id == 1:
                            filtered_df = filtered_df[filtered_df["Q_Type"] == "MCQ"]
                        elif q_id == 2:
                            filtered_df = filtered_df[filtered_df["Q_Type"] == "SQ"]
                        elif q_id == 5:
                            filtered_df = filtered_df[filtered_df["Q_Type"] == "LQ"]

                    if selected_unit and "ٹاپک وائز" in selected_mode:
                        if "Unit" in filtered_df.columns:
                            filtered_df = filtered_df[filtered_df["Unit"].astype(str).str.strip() == str(selected_unit).strip()]
                        if selected_topics and "Topic_No" in filtered_df.columns:
                            filtered_df["Topic_No_Str"] = filtered_df["Topic_No"].astype(str).str.strip()
                            filtered_df = filtered_df[filtered_df["Topic_No_Str"].isin([str(t) for t in selected_topics])]

                    selected_records_for_q = []
                    is_saved = q_id in st.session_state.saved_sections

                    for idx, row in enumerate(filtered_df.iterrows(), 1):
                        r_data = row[1]
                        q_text = str(r_data.get("Q_Text", "")) if pd.notna(r_data.get("Q_Text")) else ""
                        expr = str(r_data.get("Expression", "")) if pd.notna(r_data.get("Expression")) else ""

                        st.markdown("<div class='unique-question-card'>", unsafe_allow_html=True)
                        cols = st.columns([0.05, 0.08, 0.35, 0.52])
                        
                        with cols[3]:
                            if q_text.strip():
                                st.markdown(f"<div style='text-align: right; direction: rtl;'>{q_text}</div>", unsafe_allow_html=True)

                            if q_id == 1:
                                opt_a = str(r_data.get("Option_A", "")) if pd.notna(r_data.get("Option_A")) else ""
                                opt_b = str(r_data.get("Option_B", "")) if pd.notna(r_data.get("Option_B")) else ""
                                opt_c = str(r_data.get("Option_C", "")) if pd.notna(r_data.get("Option_C")) else ""
                                opt_d = str(r_data.get("Option_D", "")) if pd.notna(r_data.get("Option_D")) else ""
                                
                                if opt_a or opt_b or opt_c or opt_d:
                                    opt_cols = st.columns(4)
                                    with opt_cols[0]:
                                        st.markdown("<div style='text-align: right;'>(A)</div>", unsafe_allow_html=True)
                                        clean_a = opt_a.replace("$", "").strip()
                                        if clean_a:
                                            st.latex(clean_a)
                                    with opt_cols[1]:
                                        st.markdown("<div style='text-align: right;'>(B)</div>", unsafe_allow_html=True)
                                        clean_b = opt_b.replace("$", "").strip()
                                        if clean_b:
                                            st.latex(clean_b)
                                    with opt_cols[2]:
                                        st.markdown("<div style='text-align: right;'>(C)</div>", unsafe_allow_html=True)
                                        clean_c = opt_c.replace("$", "").strip()
                                        if clean_c:
                                            st.latex(clean_c)
                                    with opt_cols[3]:
                                        st.markdown("<div style='text-align: right;'>(D)</div>", unsafe_allow_html=True)
                                        clean_d = opt_d.replace("$", "").strip()
                                        if clean_d:
                                            st.latex(clean_d)

                        with cols[2]:
                            if expr.strip():
                                clean_expr = expr.strip().replace("$", "")
                                st.latex(clean_expr)

                        with cols[1]:
                            st.markdown(f"<div style='text-align: right; margin: 0px; padding: 0px;'>**({idx})**</div>", unsafe_allow_html=True)

                        with cols[0]:
                            is_selected = st.checkbox("", key=f"q_{q_id}_{idx}", disabled=is_saved)

                        st.markdown("</div>", unsafe_allow_html=True)

                        if is_selected:
                            selected_records_for_q.append(r_data)
                    
                    if st.button(f"محفوظ کریں" if not is_saved else "محفوظ ہو گیا (لاک)", key=f"save_btn_{q_id}", disabled=is_saved):
                        st.session_state.section_data_records[q_id] = selected_records_for_q
                        st.session_state.saved_sections.add(q_id)
                        
                        if q_id == 1:
                            st.session_state.selected_section_tab = "SQs"
                        elif q_id == 2:
                            st.session_state.selected_section_tab = "LQs"
                        
                        st.success(f"سیکشن کامیابی سے محفوظ اور لاک ہو گیا ہے!")
                        st.rerun()

                except Exception as e:
                    st.error(f"ڈیٹا لوڈ کرنے میں مسئلہ پیش آیا: {e}")

                if {1, 2, 5}.issubset(st.session_state.saved_sections):
                    st.markdown("---")
                    st.success("🎉 آپ کا پیپر بن چکا ہے!")

                    def generate_paper_document(sel_cls, sel_subj, sel_unit, sel_top, inst_name, logo_file):
                        doc = Document()
                        for section in doc.sections:
                            section.page_width = Inches(8.27)
                            section.page_height = Inches(11.69)
                            section.top_margin = Inches(0.4)
                            section.bottom_margin = Inches(0.4)
                            section.left_margin = Inches(0.4)
                            section.right_margin = Inches(0.4)
                            
                            body = section._sectPr
                            bidi = OxmlElement('w:bidi')
                            body.append(bidi)

                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Jameel Noori Nastaleeq'
                        font.size = Pt(12)
                        
                        rPr = style.element.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Jameel Noori Nastaleeq')
                        rFonts.set(qn('w:hAnsi'), 'Jameel Noori Nastaleeq')
                        rFonts.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
                        rPr.append(rFonts)

                        total_paper_marks = 0
                        for q_idx, records in st.session_state.section_data_records.items():
                            cnt = len(records)
                            if q_idx == 1:
                                total_paper_marks += cnt * 1
                            elif q_idx == 2:
                                total_paper_marks += cnt * 2
                            elif q_idx == 5:
                                total_paper_marks += cnt * 4

                        raw_minutes = total_paper_marks * 1.5
                        remainder = raw_minutes % 10
                        if remainder >= 5:
                            total_minutes = int(raw_minutes + (10 - remainder))
                        else:
                            total_minutes = int(raw_minutes - remainder)

                        sel_class_display = sel_cls.replace(" Class", "").strip()

                        class_urdu_map = {
                            "9th": "نہم", "10th": "دہم",
                            "8th": "آٹھویں", "7th": "ساتویں"
                        }
                        sel_class_urdu = class_urdu_map.get(sel_class_display, sel_class_display)

                        subject_urdu_map = {
                            "Chemistry": "کیمسٹری", "Physics": "فزکس", "Biology": "بائیولوجی",
                            "Mathematics": "ریاضی", "Math": "ریاضی", "English": "انگریزی", 
                            "Urdu": "اردو", "Islamiyat": "اسلامیات", "Pak_Studies": "مطالعہ پاکستان", 
                            "Computer": "کمپیوٹر سائنس"
                        }
                        sel_subj_urdu = subject_urdu_map.get(sel_subj, sel_subj)
                        sel_unit_val = selected_unit if selected_unit else "Full_Book"

                        def set_rtl_paragraph(paragraph, alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT):
                            paragraph.alignment = alignment
                            pPr = paragraph._p.get_or_add_pPr()
                            pPr.append(OxmlElement('w:bidi'))

                        def set_rtl_table(table):
                            tblPr = table._tbl.tblPr
                            tblPr.append(OxmlElement('w:bidiVisual'))

                        def append_rtl_text(paragraph, text_str):
                            if not text_str or not str(text_str).strip():
                                return
                            run = paragraph.add_run(str(text_str).strip())
                            run.font.name = 'Jameel Noori Nastaleeq'
                            run.font.size = Pt(12)
                            run.font.rtl = True
                            rPr = run._r.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Jameel Noori Nastaleeq')
                            rFonts.set(qn('w:hAnsi'), 'Jameel Noori Nastaleeq')
                            rFonts.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
                            rPr.append(rFonts)

                        head_table = doc.add_table(rows=1, cols=8)
                        head_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        head_table.autofit = False
                        set_rtl_table(head_table)

                        col_width = Inches(7.47 / 8.0)
                        for i in range(8):
                            head_table.columns[i].width = col_width

                        trPr = head_table.rows[0]._tr.get_or_add_trPr()
                        trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                        trPr.append(trHeight)

                        cells = head_table.rows[0].cells
                        cell_left = cells[0]
                        cell_mid = cells[1].merge(cells[6])
                        cell_right = cells[7]

                        cell_left.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        p_left = cell_left.paragraphs[0]
                        set_rtl_paragraph(p_left, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        p_left.paragraph_format.space_after = Pt(0)
                        p_left.paragraph_format.space_before = Pt(0)

                        if logo_file is not None:
                            temp_logo_path = "temp_uploaded_logo.png"
                            with open(temp_logo_path, "wb") as f:
                                f.write(logo_file.getbuffer())
                            p_left.add_run().add_picture(temp_logo_path, width=Inches(0.5))

                        cell_mid.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        p_mid = cell_mid.paragraphs[0]
                        set_rtl_paragraph(p_mid, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        p_mid.paragraph_format.space_after = Pt(0)
                        p_mid.paragraph_format.space_before = Pt(0)

                        final_inst_name = inst_name.strip() if inst_name.strip() else "GOVERNMENT HIGH SCHOOL KHAN PUR MARAL MULTAN"
                        run_m1 = p_mid.add_run(f"{final_inst_name.upper()}\n")
                        run_m1.font.name = 'Jameel Noori Nastaleeq'
                        run_m1.font.size = Pt(20)
                        run_m1.bold = True
                        run_m1.font.color.rgb = RGBColor(30, 58, 138)
                        run_m1.font.rtl = True

                        run_m2 = p_mid.add_run("TOPIC WISE TEST")
                        run_m2.font.name = 'Jameel Noori Nastaleeq'
                        run_m2.font.size = Pt(14)
                        run_m2.bold = True
                        run_m2.font.color.rgb = RGBColor(100, 100, 100)
                        run_m2.font.rtl = True

                        cell_right.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        p_right = cell_right.paragraphs[0]
                        set_rtl_paragraph(p_right, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                        p_right.paragraph_format.space_after = Pt(0)
                        p_right.paragraph_format.space_before = Pt(0)
                        
                        run_r = p_right.add_run(f"{sel_subj_urdu}\n({sel_class_urdu})")
                        run_r.font.name = 'Jameel Noori Nastaleeq'
                        run_r.font.size = Pt(20)
                        run_r.font.rtl = True
                        run_r.bold = True
                        run_r.font.color.rgb = RGBColor(30, 58, 138)
                        rPr_r = run_r._r.get_or_add_rPr()
                        rFonts_r = OxmlElement('w:rFonts')
                        rFonts_r.set(qn('w:ascii'), 'Jameel Noori Nastaleeq')
                        rFonts_r.set(qn('w:hAnsi'), 'Jameel Noori Nastaleeq')
                        rFonts_r.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
                        rPr_r.append(rFonts_r)

                        for c in [cell_left, cell_mid, cell_right]:
                            shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(shd)
                            
                            tcMar = parse_xml(r'''
                                <w:tcMar {} >
                                    <w:top w:w="60" w:type="dxa"/>
                                    <w:bottom w:w="60" w:type="dxa"/>
                                    <w:left w:w="60" w:type="dxa"/>
                                    <w:right w:w="60" w:type="dxa"/>
                                </w:tcMar>
                            '''.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(tcMar)

                            tcBorders = parse_xml(r'''
                                <w:tcBorders {} >
                                    <w:top w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>
                                    <w:left w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>
                                    <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>
                                    <w:right w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>
                                </w:tcBorders>
                            '''.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(tcBorders)
                        topics_str_formatted = " + ".join([str(t) for t in sel_top]) if sel_top else "All Topics"
                        unit_info_str = f"Unit: {sel_unit_val}" if sel_unit_val != "Full_Book" else "Full Book"
                        
                        info_table = doc.add_table(rows=1, cols=4)
                        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        info_table.autofit = False
                        set_rtl_table(info_table)
                        
                        info_widths = [Inches(1.4), Inches(1.5), Inches(3.2), Inches(1.4)]
                        for idx, w in enumerate(info_widths):
                            info_table.columns[idx].width = w

                        trPr_info = info_table.rows[0]._tr.get_or_add_trPr()
                        trHeight_info = parse_xml(r'<w:trHeight {} w:val="380" w:hRule="atLeast"/>'.format(nsdecls('w')))
                        trPr_info.append(trHeight_info)

                        info_cells = info_table.rows[0].cells
                        info_texts = [
                            f"Marks: {total_paper_marks}",
                            unit_info_str,
                            f"Topics: {topics_str_formatted}",
                            f"Time: {total_minutes} Mins"
                        ]
                        for idx, text in enumerate(info_texts):
                            c = info_cells[idx]
                            c.width = info_widths[idx]
                            c.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            p_c = c.paragraphs[0]
                            set_rtl_paragraph(p_c, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                            p_c.paragraph_format.space_after = Pt(0)
                            p_c.paragraph_format.space_before = Pt(0)
                            
                            run_c = c.paragraphs[0].add_run(text)
                            run_c.font.name = 'Arial'
                            run_c.font.size = Pt(9.5)
                            run_c.bold = True
                            run_c.font.color.rgb = RGBColor(15, 23, 42)
                            
                            shd_info = parse_xml(r'<w:shd {} w:fill="E2E8F0"/>'.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(shd_info)

                            tcMar_info = parse_xml(r'''
                                <w:tcMar {} >
                                    <w:top w:w="50" w:type="dxa"/>
                                    <w:bottom w:w="50" w:type="dxa"/>
                                    <w:left w:w="50" w:type="dxa"/>
                                    <w:right w:w="50" w:type="dxa"/>
                                </w:tcMar>
                            '''.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(tcMar_info)

                            tcBorders_info = parse_xml(r'''
                                <w:tcBorders {} >
                                    <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
                                    <w:left w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
                                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
                                    <w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
                                </w:tcBorders>
                            '''.format(nsdecls('w')))
                            c._tc.get_or_add_tcPr().append(tcBorders_info)

                        for q_idx in sorted(st.session_state.section_data_records.keys()):
                            if st.session_state.section_data_records[q_idx]:
                                q_count = len(st.session_state.section_data_records[q_idx])

                                p_part = doc.add_paragraph()
                                set_rtl_paragraph(p_part, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                p_part.paragraph_format.space_before = Pt(8)
                                p_part.paragraph_format.space_after = Pt(2)

                                run_part = p_part.add_run()
                                run_part.font.name = 'Jameel Noori Nastaleeq'
                                run_part.font.rtl = True
                                run_part.bold = True
                                run_part.font.size = Pt(16)
                                rPr_p = run_part._r.get_or_add_rPr()
                                rFonts_p = OxmlElement('w:rFonts')
                                rFonts_p.set(qn('w:ascii'), 'Jameel Noori Nastaleeq')
                                rFonts_p.set(qn('w:hAnsi'), 'Jameel Noori Nastaleeq')
                                rFonts_p.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
                                rPr_p.append(rFonts_p)

                                if q_idx == 1:
                                    run_part.text = "حصہ اول (معروضی)"
                                elif q_idx == 2:
                                    run_part.text = "حصہ دوم (مختصر سوالات)"
                                elif q_idx == 5:
                                    run_part.text = "حصہ سوم (طویل سوالات)"

                                p_heading = doc.add_paragraph()
                                set_rtl_paragraph(p_heading, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                p_heading.paragraph_format.space_before = Pt(4)
                                p_heading.paragraph_format.space_after = Pt(4)
                                
                                run_h = p_heading.add_run()
                                run_h.font.name = 'Jameel Noori Nastaleeq'
                                run_h.font.rtl = True  
                                run_h.bold = True
                                run_h.font.size = Pt(12)
                                rPr_h = run_h._r.get_or_add_rPr()
                                rFonts_h = OxmlElement('w:rFonts')
                                rFonts_h.set(qn('w:ascii'), 'Jameel Noori Nastaleeq')
                                rFonts_h.set(qn('w:hAnsi'), 'Jameel Noori Nastaleeq')
                                rFonts_h.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
                                rPr_h.append(rFonts_h)
                                
                                if q_idx == 1:
                                    total_m = q_count * 1
                                    run_h.text = f"سوال نمبر 1\tدرست جواب پر دائرہ لگائیں۔\tکل نمبر: {total_m}"
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
                                elif q_idx == 2:
                                    total_m = q_count * 2
                                    run_h.text = f"سوال نمبر 2\tدرج ذیل مختصر سوالات کے جوابات دیں۔\tکل نمبر: {total_m}"
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
                                elif q_idx == 5:
                                    total_m = q_count * 4
                                    run_h.text = f"سوال نمبر 3\tدرج ذیل طویل سوالات کے جوابات دیں۔\tکل نمبر: {total_m}"
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
                                    p_heading.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))

                                pBdr_h = OxmlElement('w:pBdr')
                                top_h = OxmlElement('w:top')
                                top_h.set(qn('w:val'), 'single')
                                top_h.set(qn('w:sz'), '6')
                                top_h.set(qn('w:space'), '1')
                                top_h.set(qn('w:color'), '000000')
                                pBdr_h.append(top_h)

                                bottom_h = OxmlElement('w:bottom')
                                bottom_h.set(qn('w:val'), 'single')
                                bottom_h.set(qn('w:sz'), '6')
                                bottom_h.set(qn('w:space'), '1')
                                bottom_h.set(qn('w:color'), '000000')
                                pBdr_h.append(bottom_h)

                                p_heading._p.get_or_add_pPr().append(pBdr_h)

                                if q_idx == 2:
                                    # --- SQs (مختصر سوالات) ڈائنامک ٹیبل سیٹنگ: فکسڈ 6 کالم ---
                                    records = st.session_state.section_data_records[q_idx]
                                    total_q = len(records)
                                    num_rows = (total_q + 1) // 2
                                    
                                    sq_table = doc.add_table(rows=num_rows, cols=6)
                                    sq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    sq_table.autofit = False
                                    set_rtl_table(sq_table)                                    
                                    
                                    # کالمز کی لمبائی: 0.5, 1.5, 1.5, 0.5, 1.5, 1.5 (ٹوٹل 7.0 انچ)
                                    col_widths = [Inches(0.5), Inches(1.50), Inches(1.50), Inches(0.5), Inches(1.50), Inches(1.50)]
                                    
                                    for row_i in range(num_rows):
                                        for col_j in range(6):
                                            sq_table.rows[row_i].cells[col_j].width = col_widths[col_j]
                                            sq_table.rows[row_i].cells[col_j].vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        
                                        trPr_sq = sq_table.rows[row_i]._tr.get_or_add_trPr()
                                        trHeight_sq = parse_xml(r'<w:trHeight {} w:val="350" w:hRule="atLeast"/>'.format(nsdecls('w')))
                                        trPr_sq.append(trHeight_sq)              
                                        
                                        # پہلا سوال (لیفٹ سائیڈ)
                                        idx_item_1 = row_i * 2
                                        if idx_item_1 < total_q:
                                            r_data_1 = records[idx_item_1]               
                                            
                                            # کالم 1: نمبرنگ (سنٹر)
                                            c_n1 = sq_table.rows[row_i].cells[0]
                                            p_n1 = c_n1.paragraphs[0]
                                            set_rtl_paragraph(p_n1, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                            run_n1 = p_n1.add_run(f"({idx_item_1 + 1})")
                                            run_n1.font.name = 'Jameel Noori Nastaleeq'
                                            run_n1.font.size = Pt(12)
                                            run_n1.font.rtl = True
                                            run_n1.bold = True

                                            q_text_1 = str(r_data_1.get("Q_Text", "")) if pd.notna(r_data_1.get("Q_Text")) else ""
                                            expr_1 = str(r_data_1.get("Expression", "")) if pd.notna(r_data_1.get("Expression")) else ""
                                            has_expr_1 = bool(expr_1.strip())

                                            if not has_expr_1:
                                                # اگر ایکسپریشن نہیں ہے تو کالم 2 اور 3 کو مرج کر دیں
                                                merged_cell_1 = sq_table.rows[row_i].cells[1].merge(sq_table.rows[row_i].cells[2])
                                                p_t1 = merged_cell_1.paragraphs[0]
                                                set_rtl_paragraph(p_t1, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                                if q_text_1:
                                                    append_rtl_text(p_t1, q_text_1)
                                            else:
                                                # کالم 2: ٹیکسٹ (لیفٹ الائن)
                                                c_t1 = sq_table.rows[row_i].cells[1]
                                                p_t1 = c_t1.paragraphs[0]
                                                set_rtl_paragraph(p_t1, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                                if q_text_1:
                                                    append_rtl_text(p_t1, q_text_1)

                                                # کالم 3: ایکسپریشن (سنٹر الائن)
                                                c_e1 = sq_table.rows[row_i].cells[2]
                                                p_e1 = c_e1.paragraphs[0]
                                                set_rtl_paragraph(p_e1, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                                add_math_expression_as_image(p_e1, expr_1, font_size=28)

                                        # دوسرا سوال (رائٹ سائیڈ)
                                        idx_item_2 = row_i * 2 + 1
                                        if idx_item_2 < total_q:
                                            r_data_2 = records[idx_item_2]
                                            
                                            # کالم 4: نمبرنگ (سنٹر)
                                            c_n2 = sq_table.rows[row_i].cells[3]
                                            p_n2 = c_n2.paragraphs[0]
                                            set_rtl_paragraph(p_n2, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                            run_n2 = p_n2.add_run(f"({idx_item_2 + 1})")
                                            run_n2.font.name = 'Jameel Noori Nastaleeq'
                                            run_n2.font.size = Pt(12)
                                            run_n2.font.rtl = True
                                            run_n2.bold = True

                                            q_text_2 = str(r_data_2.get("Q_Text", "")) if pd.notna(r_data_2.get("Q_Text")) else ""
                                            expr_2 = str(r_data_2.get("Expression", "")) if pd.notna(r_data_2.get("Expression")) else ""
                                            has_expr_2 = bool(expr_2.strip())

                                            if not has_expr_2:
                                                # اگر ایکسپریشن نہیں ہے تو کالم 5 اور 6 کو مرج کر دیں
                                                merged_cell_2 = sq_table.rows[row_i].cells[4].merge(sq_table.rows[row_i].cells[5])
                                                p_t2 = merged_cell_2.paragraphs[0]
                                                set_rtl_paragraph(p_t2, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                                if q_text_2:
                                                    append_rtl_text(p_t2, q_text_2)
                                            else:
                                                # کالم 5: ٹیکسٹ (لیفٹ الائن)
                                                c_t2 = sq_table.rows[row_i].cells[4]
                                                p_t2 = c_t2.paragraphs[0]
                                                set_rtl_paragraph(p_t2, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                                if q_text_2:
                                                    append_rtl_text(p_t2, q_text_2)

                                                # کالم 6: ایکسپریشن (سنٹر الائن)
                                                c_e2 = sq_table.rows[row_i].cells[5]
                                                p_e2 = c_e2.paragraphs[0]
                                                set_rtl_paragraph(p_e2, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                                add_math_expression_as_image(p_e2, expr_2, font_size=28)

                                elif q_idx == 5:
                                    # --- LQs (طویل سوالات) ڈائنامک ٹیبل سیٹنگ ---
                                    records = st.session_state.section_data_records[q_idx]
                                    for sub_i, r_data in enumerate(records, 1):
                                        lq_table = doc.add_table(rows=1, cols=5)
                                        lq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                        lq_table.autofit = False
                                        set_rtl_table(lq_table)

                                        # نئے کالمز کی چوڑائی: 1.0, 0.5, 3.5, 1.5, 0.5 (ٹوٹل 7.0 انچ)
                                        lq_widths = [Inches(1.0), Inches(0.5), Inches(3.5), Inches(1.5), Inches(0.5)]
                                        for i, w in enumerate(lq_widths):
                                            lq_table.columns[i].width = w

                                        trPr_lq = lq_table.rows[0]._tr.get_or_add_trPr()
                                        trHeight_lq = parse_xml(r'<w:trHeight {} w:val="380" w:hRule="atLeast"/>'.format(nsdecls('w')))
                                        trPr_lq.append(trHeight_lq)

                                        cells_lq = lq_table.rows[0].cells
                                        for i, w in enumerate(lq_widths):
                                            cells_lq[i].width = w
                                            cells_lq[i].vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER

                                        # کالم 1: سوال نمبر (سنٹر الائن)
                                        c_num = cells_lq[0]
                                        p_num = c_num.paragraphs[0]
                                        set_rtl_paragraph(p_num, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                        run_num = p_num.add_run(f"سوال نمبر {sub_i + 2}")
                                        run_num.font.name = 'Jameel Noori Nastaleeq'
                                        run_num.font.size = Pt(12)
                                        run_num.font.rtl = True
                                        run_num.bold = True

                                        # کالم 2: جزو A یا B (سنٹر الائن)
                                        c_sub = cells_lq[1]
                                        p_sub = c_sub.paragraphs[0]
                                        set_rtl_paragraph(p_sub, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                        sub_part_val = str(r_data.get("Sub_Part", "A")).strip().upper()
                                        if sub_part_val not in ["A", "B"]:
                                            sub_part_val = "A"
                                        run_sub = p_sub.add_run(f"({sub_part_val})")
                                        run_sub.font.name = 'Jameel Noori Nastaleeq'
                                        run_sub.font.size = Pt(12)
                                        run_sub.font.rtl = True
                                        run_sub.bold = True

                                        marks_val = "4"
                                        expr = r_data.get("Expression", "")
                                        has_expr = pd.notna(expr) and str(expr).strip() != ""

                                        if not has_expr:
                                            # اگر ایکسپریشن نہیں ہے تو کالم 3 اور 4 مرج کر دیں اور لیفٹ الائن کریں
                                            merged_cell_34 = cells_lq[2].merge(cells_lq[3])
                                            p_text = merged_cell_34.paragraphs[0]
                                            set_rtl_paragraph(p_text, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                            q_text_val = str(r_data.get("Q_Text", "")) if pd.notna(r_data.get("Q_Text")) else ""
                                            if q_text_val:
                                                append_rtl_text(p_text, q_text_val)
                                        else:
                                            # کالم 3: سوال کی ٹیکسٹ (سنٹر الائن)
                                            c_txt = cells_lq[2]
                                            p_txt = c_txt.paragraphs[0]
                                            set_rtl_paragraph(p_txt, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                            q_text_val = str(r_data.get("Q_Text", "")) if pd.notna(r_data.get("Q_Text")) else ""
                                            if q_text_val:
                                                append_rtl_text(p_txt, q_text_val)

                                            # کالم 4: ایکسپریشن (لیفٹ الائن)
                                            c_expr = cells_lq[3]
                                            p_expr = c_expr.paragraphs[0]
                                            set_rtl_paragraph(p_expr, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                            add_math_expression_as_image(p_expr, str(expr), font_size=28)

                                        # کالم 5: نمبر 4 (سنٹر الائن)
                                        c_marks = cells_lq[4]
                                        p_marks = c_marks.paragraphs[0]
                                        set_rtl_paragraph(p_marks, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                        run_marks = p_marks.add_run(marks_val)
                                        run_marks.font.name = 'Jameel Noori Nastaleeq'
                                        run_marks.font.size = Pt(12)
                                        run_marks.font.rtl = True
                                        run_marks.bold = True

                                else:
                                    for sub_i, r_data in enumerate(st.session_state.section_data_records[q_idx], 1):
                                        # --- پہلی رو: سوال (تین کالم) ---
                                        mcq_table = doc.add_table(rows=1, cols=3)
                                        mcq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                        mcq_table.autofit = False
                                        set_rtl_table(mcq_table)

                                        # کالم کی لمبائی: 0.75، 1.75، 4.50 انچ
                                        col_widths_q = [Inches(0.75), Inches(1.75), Inches(4.50)]
                                        for i, width in enumerate(col_widths_q):
                                            mcq_table.columns[i].width = width

                                        trPr_mcq = mcq_table.rows[0]._tr.get_or_add_trPr()
                                        trHeight_mcq = parse_xml(r'<w:trHeight {} w:val="350" w:hRule="atLeast"/>'.format(nsdecls('w')))
                                        trPr_mcq.append(trHeight_mcq)

                                        cells_q = mcq_table.rows[0].cells
                                        for i, width in enumerate(col_widths_q):
                                            cells_q[i].width = width
                                            cells_q[i].vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER

                                        # پہلا کالم: نمبرنگ (سنٹر)
                                        p_num = cells_q[0].paragraphs[0]
                                        set_rtl_paragraph(p_num, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                        run_num = p_num.add_run(f"({sub_i})")
                                        run_num.font.name = 'Jameel Noori Nastaleeq'
                                        run_num.font.size = Pt(12)
                                        run_num.font.rtl = True
                                        run_num.bold = True

                                        expr = r_data.get("Expression", "")
                                        q_text_val = str(r_data.get("Q_Text", "")) if pd.notna(r_data.get("Q_Text")) else ""
                                        has_expr = pd.notna(expr) and str(expr).strip() != ""

                                        if not has_expr:
                                            # اگر ایکسپریشن نہیں ہے تو دوسرے اور تیسرے کالم کو مرج کر دیں
                                            merged_cell = cells_q[1].merge(cells_q[2])
                                            p_txt = merged_cell.paragraphs[0]
                                            set_rtl_paragraph(p_txt, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                            if q_text_val:
                                                append_rtl_text(p_txt, q_text_val)
                                        else:
                                            # دوسرا کالم: ایکسپریشن (سنٹر)
                                            p_expr = cells_q[1].paragraphs[0]
                                            set_rtl_paragraph(p_expr, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                            add_math_expression_as_image(p_expr, str(expr), font_size=28)

                                            # تیسرا کالم: سوال کی ٹیکسٹ (لیفٹ الائن)
                                            p_txt = cells_q[2].paragraphs[0]
                                            set_rtl_paragraph(p_txt, docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT)
                                            if q_text_val:
                                                append_rtl_text(p_txt, q_text_val)

                                        # --- دوسری رو: آپشنز (آٹھ کالم) ---
                                        opt_a = str(r_data.get("Option_A", "")) if pd.notna(r_data.get("Option_A")) else ""
                                        opt_b = str(r_data.get("Option_B", "")) if pd.notna(r_data.get("Option_B")) else ""
                                        opt_c = str(r_data.get("Option_C", "")) if pd.notna(r_data.get("Option_C")) else ""
                                        opt_d = str(r_data.get("Option_D", "")) if pd.notna(r_data.get("Option_D")) else ""

                                        if opt_a or opt_b or opt_c or opt_d:
                                            opt_table = doc.add_table(rows=1, cols=8)
                                            opt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                            opt_table.autofit = False
                                            set_rtl_table(opt_table)
                                            
                                            # کالمز کی لمبائی: 0.5, 1.25, 0.5, 1.25, 0.5, 1.25, 0.5, 1.25
                                            col_widths_opt = [Inches(0.5), Inches(1.25), Inches(0.5), Inches(1.25), 
                                                              Inches(0.5), Inches(1.25), Inches(0.5), Inches(1.25)]
                                            for i, width in enumerate(col_widths_opt):
                                                opt_table.columns[i].width = width

                                            hdr_cells = opt_table.rows[0].cells
                                            for i, width in enumerate(col_widths_opt):
                                                hdr_cells[i].width = width
                                                hdr_cells[i].vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER

                                            options_data = [("(A)", opt_a), ("(B)", opt_b), ("(C)", opt_c), ("(D)", opt_d)]
                                            
                                            for idx, (label, val) in enumerate(options_data):
                                                cell_lbl = hdr_cells[idx * 2]
                                                cell_val = hdr_cells[idx * 2 + 1]
                                                
                                                # لیبل A, B, C, D (سنٹر)
                                                p_lbl = cell_lbl.paragraphs[0]
                                                set_rtl_paragraph(p_lbl, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                                run_lbl = p_lbl.add_run(label)
                                                run_lbl.font.name = 'Jameel Noori Nastaleeq'
                                                run_lbl.font.size = Pt(12)
                                                run_lbl.font.rtl = True
                                                
                                                # آپشن کی ویلیو (سنٹر)
                                                p_val = cell_val.paragraphs[0]
                                                set_rtl_paragraph(p_val, docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
                                                
                                                if val and str(val).strip():
                                                    val_str = str(val).strip()
                                                    if "\\" in val_str or "^" in val_str or "_" in val_str or "$" in val_str:
                                                        add_math_expression_as_image(p_val, val_str, font_size=28)
                                                    else:
                                                        append_rtl_text(p_val, val_str)

                        return doc

                    target_folder = r"F:\Automated_Paper_App"
                    os.makedirs(target_folder, exist_ok=True)

                    if selected_topics:
                        file_naming_part = f"Unit_{selected_unit}_Topics_Selected"
                    else:
                        file_naming_part = f"Unit_{selected_unit}"

                    for char in ['\\', '/', ':', '*', '?', '"', '<', '>', '|', '+']:
                        file_naming_part = file_naming_part.replace(char, '_')

                    file_name = f"{selected_cls.replace(' Class', '').strip()}_{selected_subj}_{file_naming_part}".replace(" ", "_")
                    file_path_docx = os.path.join(target_folder, f"{file_name}.docx")

                    col_b1, _ = st.columns(2)
                    
                    with col_b1:
                        if st.button("ورڈ پیپر"):
                            if not st.session_state.section_data_records:
                                st.warning("براہ کرم پہلے کوئی سیکشن محفوظ کریں!")
                            else:
                                try:
                                    doc = generate_paper_document(selected_cls, selected_subj, selected_unit, selected_topics, institution_name, uploaded_logo)
                                    doc.save(file_path_docx)
                                    with open(file_path_docx, "rb") as f_docx:
                                        st.download_button(
                                            label="ورڈ فائل ڈاؤن لوڈ کریں",
                                            data=f_docx,
                                            file_name=f"{file_name}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )
                                except Exception as e:
                                    st.error(f"مسئلہ پیش آیا: {e}")